"""Generate summarized InternTrack analysis Word document (aligned to codebase)."""
from pathlib import Path
from datetime import date

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    import subprocess, sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT


def set_run(run, bold=False, size=11, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, bold=bold, size=size)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_run(r, bold=True, size=10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val))
            set_run(r, size=10)
    doc.add_paragraph()
    return table


def main():
    today = date(2026, 7, 21)
    out = Path(__file__).resolve().parent / f"INTERNTRACK_Analysis_Summary_{today.isoformat()}.docx"
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("INTERNTRACK")
    set_run(r, bold=True, size=22)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Complete System Analysis — Summary Report")
    set_run(r, bold=True, size=14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        "University of Cabuyao (Pamantasan ng Cabuyao)\n"
        "Evidence-based review vs Internship Manual 2023, Capstone Ch.1–3, Adviser/Client Transcript\n"
        f"Date: 21 July 2026 (Reverb + chat + meetings + e-sign in product scope)"
    )
    set_run(r, size=10, color=(80, 80, 80))

    doc.add_paragraph()

    add_heading_styled(doc, "1. Overall Project Percentage", 1)
    add_para(doc, "OVERALL PROJECT COMPLETION: 86%", bold=True, size=16)
    add_para(
        doc,
        "Weighted from implemented workflows and Manual/adviser compliance. Uplifted after Laravel Reverb "
        "live notifications (Echo + poll fallback), internship-scoped messaging, meetings/RSVP, and canvas "
        "electronic signatures on evaluations and document approvals.",
    )

    add_table(
        doc,
        ["Metric", "Percentage / Status"],
        [
            ["Overall Project Completion", "86%"],
            ["Objective 1 — Design & Develop System", "93%"],
            ["Scope of the Study Completion", "85%"],
            ["Internship Manual Compliance", "76%"],
            ["Adviser / Client Requirement Compliance", "88%"],
            ["Objectives 2–3 — ISO Evaluation (survey)", "Not in codebase (research phase — instruments outside app)"],
            ["Production Deployment Readiness", "Not Ready (hours policy + live MISD pending)"],
            ["Capstone Defense Readiness", "Near-Ready"],
        ],
    )

    add_heading_styled(doc, "2. Executive Summary", 1)
    add_para(
        doc,
        "INTERNTRACK is a working multi-role internship management system (React + Laravel Sanctum + MySQL) "
        "for the University of Cabuyao. Portals: Student, Faculty, Coordinator, PALD Director, Industry Supervisor, "
        "MISD Admin. Modules include documents, weekly FO-31 journals, attendance, placement/MOA gate, portfolio "
        "(active), evaluations with e-sign, Director absorption, internship messaging, meetings/orientation RSVP, "
        "and live notifications via Reverb.",
    )
    add_para(
        doc,
        "Defense wording that must match the code: (1) Journals = weekly upload of the FO-31 Daily Journal form; "
        "(2) Realtime = Reverb WebSockets when VITE_REVERB_APP_KEY is set, else ~60s poll fallback; "
        "(3) MISD = mock/local + Admin Sync, not live iEnroll; "
        "(4) E-sign = drawn PNG + typed name + timestamp (acknowledgment, not PKI).",
        bold=True,
    )

    add_heading_styled(doc, "3. System Overview (Brief)", 1)
    add_para(
        doc,
        "Stack: React 18 + Vite · Laravel 12 REST API · MySQL · Sanctum · Laravel Reverb (optional for demos). "
        "Run API + `php artisan reverb:start` + Vite for live inbox/chat updates.",
    )

    add_heading_styled(doc, "4. Business Process Status Summary", 1)
    add_table(
        doc,
        ["Status", "Count", "Share"],
        [
            ["Fully Implemented", "20", "69%"],
            ["Partially Implemented", "5", "17%"],
            ["Incorrectly Implemented", "1", "3%"],
            ["Missing / Phase 2", "3", "10%"],
        ],
    )
    add_para(doc, "Key Fully Done:", bold=True)
    add_para(
        doc,
        "MISD Admin + mock sync · MOA/slot hard-gate · Status timeline + ownership · Document routing · "
        "13 required docs · Weekly FO-31 journals · ePortfolio (active) · Prefs · Reverb + Echo (poll fallback) · "
        "Internship chat · Meetings/RSVP · Canvas e-sign · Director absorption · Faculty/Coord/Director reports · At-risk SQL",
    )
    add_para(doc, "Phase 2 / future (do not claim as shipped):", bold=True)
    add_para(
        doc,
        "Browser Web Push / FCM · Chat attachments · PKI/DigiSign · Recurring meetings/ICS · Live institutional MISD · Archive + 3-year trends",
    )

    add_heading_styled(doc, "5. Strengths & Weaknesses", 1)
    add_para(doc, "Strengths", bold=True)
    add_para(
        doc,
        "• Multi-role portals including MISD Admin\n"
        "• Document routing with e-sign on approve/verify\n"
        "• Live notifications (Reverb) with honest poll fallback\n"
        "• Internship-scoped messaging and orientation meetings\n"
        "• Director-owned absorption / hire finalization\n"
        "• Hard MOA eligibility + slot inventory",
    )
    add_para(doc, "Weaknesses", bold=True)
    add_para(
        doc,
        "• Hours policy mismatch vs Internship Manual (360 default)\n"
        "• Live institutional MISD not connected (mock only)\n"
        "• E-sign is acknowledgment-grade, not qualified digital signature\n"
        "• Security depth: limited Policies; no real MFA",
    )

    add_heading_styled(doc, "6. Priority Recommendations (Summarized)", 1)
    add_table(
        doc,
        ["Priority", "Action", "Scope"],
        [
            ["DONE", "Reverb + Echo live notifications (+ poll fallback)", "CLOSED"],
            ["DONE", "Internship chat + meetings/RSVP + canvas e-sign", "CLOSED"],
            ["DONE", "MISD Admin; Director absorption; prefs; MOA gate", "CLOSED"],
            ["Critical", "Map program → target hours (Manual table)", "REQUIRED"],
            ["High", "ISO/IEC 25010 survey instruments + response analysis (Obj 2–3)", "RESEARCH — see thesis/iso25010/"],
            ["High", "Archive inactive students; certificate gates", "REQUIRED"],
            ["Phase 2", "Web Push; chat attachments; PKI; live MISD SSO", "FUTURE"],
        ],
    )

    add_heading_styled(doc, "7. Chapters 1–3 / Defense Script Alignment", 1)
    add_para(
        doc,
        "Use DEFENSE_SCRIPT.md and PROGRESS_NOTES.md:\n"
        "• Journals: weekly FO-31 Daily Journal form uploads.\n"
        "• Realtime: Reverb when VITE_REVERB_APP_KEY set; else ~60s poll — say both.\n"
        "• MISD: mock/local + Admin Sync; live iEnroll SSO considered only.\n"
        "• Director: meetings/announcements/reports — no Messages inbox.\n"
        "• Obj 2–3: external ISO instruments under thesis/iso25010/ (not an app module).",
    )

    add_heading_styled(doc, "8. Suggested Roadmap", 1)
    add_para(
        doc,
        "Done: Reverb/chat/meetings/e-sign + harden pass (MOA, prefs, absorption, reports).\n"
        "Next: Manual hours mapping, archive, certificate gates, ISO survey collection.\n"
        "Post-capstone: Live MISD, Web Push, PKI-grade signatures if required.",
    )

    add_heading_styled(doc, "9. Capstone Readiness Conclusion", 1)
    add_para(
        doc,
        "INTERNTRACK is defendable if oral and manuscript wording match the code (weekly journals, "
        "Reverb-or-poll realtime, mock MISD, chat/meetings/e-sign MVPs, Director absorption). "
        "Close Manual hours and run ISO instruments for Obj 2–3. Overall completion ~86%.",
        bold=True,
    )

    footer = doc.add_paragraph()
    r = footer.add_run(
        "\nSource: codebase CAPSTONE adad (July 2026) + Manual + PROGRESS_NOTES.md. "
        "Regenerate: python generate_analysis_summary_docx.py"
    )
    set_run(r, size=9, color=(100, 100, 100))

    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
